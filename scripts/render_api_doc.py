from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from render_solution_doc import add_inline, add_toc, configure_document, render_markdown


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "AI审核判读一致性平台-接口文档.md"
OUTPUT = ROOT.parent / "AI审核判读一致性平台-接口文档.docx"


def build() -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "AI审核判读一致性平台接口文档"
    document.core_properties.subject = "业务接口、内部接口、AI接口、Flowable回调与消息契约"
    header = document.sections[0].header.paragraphs[0]
    header.text = "AI审核判读一致性平台｜接口文档"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor(120, 130, 145)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(90)
    run = title.add_run("AI审核判读一致性平台")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor.from_string("17365D")

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(25)
    run = subtitle.add_run("接口文档")
    run.bold = True
    run.font.size = Pt(21)
    run.font.color.rgb = RGBColor.from_string("2F6F75")
    run = subtitle.add_run("\n\n用户接口｜内部接口｜AI接口｜Flowable回调｜领域事件")
    run.font.size = Pt(11.5)
    run.font.color.rgb = RGBColor(90, 100, 110)

    metadata = document.add_paragraph()
    metadata.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata.paragraph_format.space_before = Pt(95)
    metadata.paragraph_format.first_line_indent = Pt(0)
    add_inline(
        metadata,
        "文档版本：V1.0\n编制日期：2026年7月21日\n适用项目：ai-audit-calibration\n接口基线：当前工作区实际代码",
    )
    for metadata_run in metadata.runs:
        metadata_run.font.size = Pt(11)

    document.add_page_break()
    toc_heading = document.add_heading("目录", level=1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = document.add_paragraph()
    toc.paragraph_format.first_line_indent = Pt(0)
    add_toc(toc)
    document.add_page_break()

    render_markdown(document, SOURCE.read_text(encoding="utf-8"), {})
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
