from __future__ import annotations

import re
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "基于AI的审核判读一致性校准与标准化管控-项目实际落地方案.md"
OUTPUT = ROOT.parent / "基于AI的审核判读一致性校准与标准化管控-项目实际落地方案.docx"
ASSETS = ROOT / "docs" / ".generated-doc-assets"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_keep(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.keep_together = True


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("第 ")
    run.font.size = Pt(9)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)
    run = paragraph.add_run(" 页")
    run.font.size = Pt(9)


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开文档后右键更新目录"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, placeholder, end])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.15)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.first_line_indent = Pt(21)

    for name, size, color, before, after in [
        ("Heading 1", 18, "17365D", 18, 8),
        ("Heading 2", 14, "245B83", 14, 6),
        ("Heading 3", 11.5, "2F6F75", 10, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    doc.core_properties.title = "基于AI的审核判读一致性校准与标准化管控-项目实际落地方案"
    doc.core_properties.subject = "方案设计第三章四大块完整输出"
    doc.core_properties.author = "AI审核判读一致性校准项目组"

    header = section.header.paragraphs[0]
    header.text = "基于AI的审核判读一致性校准与标准化管控｜项目实际落地方案"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor(120, 130, 145)
    add_page_number(section.footer.paragraphs[0])


FONT_PATH = Path(r"C:\Windows\Fonts\msyh.ttc")


def pil_font(size: int, bold: bool = False):
    candidate = Path(r"C:\Windows\Fonts\msyhbd.ttc") if bold else FONT_PATH
    return ImageFont.truetype(str(candidate if candidate.exists() else FONT_PATH), size)


def make_canvas(title: str, height: int = 820):
    image = Image.new("RGB", (1800, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((900, 44), title, font=pil_font(44, True), fill="#17365D", anchor="mm")
    return image, draw


def pil_box(draw, box, text: str, fill: str, font_size: int = 27) -> None:
    draw.rounded_rectangle(box, radius=16, fill=fill, outline="#B7C3D0", width=3)
    draw.multiline_text(
        ((box[0] + box[2]) / 2, (box[1] + box[3]) / 2), text,
        font=pil_font(font_size), fill="#203040", anchor="mm", align="center", spacing=7
    )


def pil_arrow(draw, start, end) -> None:
    draw.line([start, end], fill="#60758A", width=5)
    x1, y1 = start; x2, y2 = end
    import math
    angle = math.atan2(y2 - y1, x2 - x1)
    length = 18
    points = [
        (x2, y2),
        (x2 - length * math.cos(angle - 0.55), y2 - length * math.sin(angle - 0.55)),
        (x2 - length * math.cos(angle + 0.55), y2 - length * math.sin(angle + 0.55)),
    ]
    draw.polygon(points, fill="#60758A")


def save_diagram(image: Image.Image, name: str) -> Path:
    path = ASSETS / name
    image.save(path, format="PNG", optimize=True)
    return path


def build_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    outputs: dict[str, Path] = {}

    image, draw = make_canvas("项目实际分层架构", 920)
    layers = [
        (135, "交互层：Vue 3 / Element Plus / ECharts / RBAC 路由", "#E8F1FA"),
        (275, "接入安全层：Gateway / JWT 双 Token / Spring Security / Sentinel / Nacos", "#EDF6F2"),
        (415, "业务域：导入｜样本｜信号｜比较｜预警｜二级真值｜统计｜风控｜学习｜追溯", "#FFF3DD"),
        (555, "流程消息层：Flowable 人工状态机｜RabbitMQ Outbox｜XXL-JOB｜Redisson", "#F6EBF7"),
        (695, "AI 与数据层：FastAPI + ONNX｜MySQL｜MongoDB｜Redis｜MinIO｜Elasticsearch", "#EAF2E5"),
    ]
    for y, text, color in layers:
        pil_box(draw, (110, y, 1690, y + 95), text, color, 29)
    for y in [230, 370, 510, 650]:
        pil_arrow(draw, (900, y), (900, y + 35))
    outputs["2.1 实际技术架构"] = save_diagram(image, "architecture.png")

    image, draw = make_canvas("审核判读闭环", 840)
    boxes = [
        ((30,180,220,285),"Excel\n导入"), ((275,180,465,285),"AI\n判读"),
        ((520,180,735,285),"一级独立\n审核"), ((790,180,1040,285),"三方逐靶标\n比较"),
        ((1130,325,1390,430),"不一致\n预警 + 二级终审"),
        ((1480,180,1765,285),"ARCHIVED\n唯一真值"),
    ]
    for box, text in boxes: pil_box(draw, box, text, "#F7FAFC", 27)
    for start, end in [((220,233),(275,233)),((465,233),(520,233)),((735,233),(790,233)),
                       ((1040,220),(1480,220)),((1040,258),(1130,365)),((1390,365),(1480,265))]:
        pil_arrow(draw, start, end)
    draw.text((1260,176), "三方一致：共同结果即真值，直接归档",
              font=pil_font(25), fill="#2F6F75", anchor="mm")
    draw.text((1095,318), "三方不一致",
              font=pil_font(24), fill="#8A5A20", anchor="mm")
    lower = [((230,575,540,680),"AI错判\n夜间增量训练","#E8F1FA"),
             ((745,575,1055,680),"一级错判\n风险培训考试","#FFF3DD"),
             ((1260,575,1645,680),"SYSTEM / PRIMARY / AI\n准确率与追溯","#EDF6F2")]
    for box,text,color in lower: pil_box(draw,box,text,color,26)
    draw.line([(1600,285),(1600,510),(385,510)], fill="#60758A", width=5)
    draw.line([(900,510),(1450,510)], fill="#60758A", width=5)
    for target in [(385,575),(900,575),(1450,575)]: pil_arrow(draw,(target[0],510),target)
    outputs["1.7 核心业务闭环"] = save_diagram(image, "workflow.png")

    image, draw = make_canvas("AI 推理输入、模型与解释输出", 780)
    items = [
        ((30,185,270,300),"原始曲线\n45点","#EAF2F8"),
        ((315,185,555,300),"清洗 / QC\nCt -100→null","#EAF2F8"),
        ((600,185,880,300),"ONNX内预处理\n前10点基线校正","#EAF2F8"),
        ((925,185,1165,300),"主曲线模型\n4类概率","#EAF2F8"),
        ((1210,185,1475,300),"结构化辅助\nCt / 浓度 / 风险","#FFF3DD"),
        ((1520,185,1770,300),"融合与后处理\n低置信→可疑","#EAF2F8"),
    ]
    for i,(box,text,color) in enumerate(items):
        pil_box(draw,box,text,color,25)
        if i < len(items)-1: pil_arrow(draw,(box[2],242),(items[i+1][0][0],242))
    pil_box(draw,(370,500,1430,630),"输出：标签 + 置信度 + 原因码 + 推理逻辑 + 特征 + 模型版本 + 降级原因","#EDF6F2",28)
    pil_arrow(draw,(1645,300),(1370,500))
    outputs["3.2 AI 推理输入输出链路"] = save_diagram(image, "ai_pipeline.png")

    image, draw = make_canvas("AI 错判反馈与持续学习闭环", 850)
    nodes = [
        ((690,120,1110,220),"归档唯一真值"), ((1280,265,1690,365),"筛选 AI 错判"),
        ((1280,525,1690,625),"夜间定时训练"), ((690,670,1110,770),"MinIO + 模型注册"),
        ((110,525,520,625),"激活 / 灰度 / 回滚"), ((110,265,520,365),"未来推理\n历史不回写"),
    ]
    for box,text in nodes: pil_box(draw,box,text,"#F3F7FA",27)
    for start,end in [((1110,170),(1280,315)),((1485,365),(1485,525)),((1280,575),(1110,720)),
                      ((690,720),(520,575)),((315,525),(315,365)),((520,315),(690,170))]: pil_arrow(draw,start,end)
    draw.multiline_text((900,425),"feedbackKey / trainingKey 幂等\n原始曲线、Ct、浓度、风险、源模型、真值完整留存",
                        font=pil_font(28),fill="#2F5D62",anchor="mm",align="center",spacing=10)
    outputs["3.8 错误样本反馈至 AI 迭代闭环"] = save_diagram(image, "feedback_loop.png")
    return outputs


INLINE_RE = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*)")


def add_inline(paragraph, text: str) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            paragraph.add_run(text[cursor:match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(42, 78, 110)
        else:
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        cursor = match.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    for row_index, values in enumerate(rows):
        row = table.rows[row_index]
        if row_index == 0:
            set_repeat_table_header(row)
        for col_index in range(width):
            cell = row.cells[col_index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, values[col_index] if col_index < len(values) else "")
            for run in paragraph.runs:
                run.font.size = Pt(8.3)
                if row_index == 0:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
            if row_index == 0:
                set_cell_shading(cell, "245B83")
            elif row_index % 2 == 0:
                set_cell_shading(cell, "F5F8FB")
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def render_markdown(doc: Document, text: str, diagrams: dict[str, Path]) -> None:
    lines = text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("# 0."))
    index = start
    in_code = False
    code_lines: list[str] = []
    while index < len(lines):
        line = lines[index].rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code:
                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.first_line_indent = Pt(0)
                paragraph.paragraph_format.left_indent = Cm(0.45)
                paragraph.paragraph_format.right_indent = Cm(0.25)
                paragraph.paragraph_format.space_before = Pt(3)
                paragraph.paragraph_format.space_after = Pt(7)
                set_cell = OxmlElement("w:shd")
                set_cell.set(qn("w:fill"), "F2F4F7")
                paragraph._p.get_or_add_pPr().append(set_cell)
                run = paragraph.add_run("\n".join(code_lines))
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
                in_code = False
                code_lines = []
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped or stripped == "---":
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[index + 1].strip().lstrip("|")):
            table_rows: list[list[str]] = []
            table_rows.append([cell.strip() for cell in stripped.strip("|").split("|")])
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_rows.append([cell.strip() for cell in lines[index].strip().strip("|").split("|")])
                index += 1
            add_table(doc, table_rows)
            continue
        heading = re.match(r"^(#{1,3})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1))
            title = heading.group(2)
            paragraph = doc.add_heading(title, level=level)
            set_repeat_keep(paragraph)
            if title in diagrams:
                picture = doc.add_picture(str(diagrams[title]), width=Inches(6.55))
                picture.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption = doc.add_paragraph(f"图：{title}")
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.paragraph_format.first_line_indent = Pt(0)
                caption.runs[0].italic = True
                caption.runs[0].font.size = Pt(8.5)
                caption.runs[0].font.color.rgb = RGBColor(110, 120, 130)
            index += 1
            continue
        if stripped.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.right_indent = Cm(0.3)
            paragraph.paragraph_format.space_before = Pt(5)
            paragraph.paragraph_format.space_after = Pt(8)
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "FFF5E6")
            paragraph._p.get_or_add_pPr().append(shd)
            add_inline(paragraph, stripped[2:])
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if bullet or numbered:
            style = "List Bullet" if bullet else "List Number"
            paragraph = doc.add_paragraph(style=style)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            add_inline(paragraph, (bullet or numbered).group(1))
            index += 1
            continue
        paragraph = doc.add_paragraph()
        add_inline(paragraph, stripped)
        index += 1


def build_document() -> None:
    source_text = SOURCE.read_text(encoding="utf-8")
    diagrams = build_diagrams()
    doc = Document()
    configure_document(doc)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover.paragraph_format.space_before = Pt(75)
    run = cover.add_run("基于AI的审核判读一致性校准\n与标准化管控")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(27)
    run.font.color.rgb = RGBColor.from_string("17365D")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(24)
    run = subtitle.add_run("项目实际落地方案")
    run.bold = True
    run.font.size = Pt(19)
    run.font.color.rgb = RGBColor.from_string("2F6F75")
    sub = subtitle.add_run("\n\n对应《方案设计》第三章四大块完整输出")
    sub.font.size = Pt(12)
    sub.font.color.rgb = RGBColor(90, 100, 110)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_before = Pt(90)
    meta.paragraph_format.first_line_indent = Pt(0)
    add_inline(meta, "文档版本：V1.0\n编制日期：2026年7月21日\n适用项目：ai-audit-calibration\n文件级别：项目设计与验收资料")
    for run in meta.runs:
        run.font.size = Pt(11)

    doc.add_page_break()
    heading = doc.add_heading("目录", level=1)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Pt(0)
    add_toc(toc)
    doc.add_page_break()

    render_markdown(doc, source_text, diagrams)

    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.1)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.15)
    section.header.is_linked_to_previous = True
    section.footer.is_linked_to_previous = True

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUTPUT)
        saved = OUTPUT
    except PermissionError:
        saved = OUTPUT.with_name(OUTPUT.stem + "-修订版" + OUTPUT.suffix)
        doc.save(saved)
    print(saved)


if __name__ == "__main__":
    build_document()
